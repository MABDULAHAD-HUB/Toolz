import os

from tkinter import *

from tkinter import filedialog, messagebox

from PIL import Image, ImageTk

from pdfminer.high_level import extract_text

from docx import Document

from fpdf import FPDF

from pillow_heif import register_heif_opener



# Register HEIF support with PIL

register_heif_opener()



# Define a class for PDF

class PDF(FPDF):

    def header(self):

        self.set_font("Arial", "B", 12)

        self.cell(0, 10, "Document Title", ln=True, align="C")



    def footer(self):

        self.set_y(-15)

        self.set_font("Arial", "I", 8)

        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C")



def sanitize_text(text):

    """Sanitize text to remove or replace unsupported characters."""

    return ''.join(c if ord(c) < 256 else ' ' for c in text)  # Replace with a space



# Conversion Functions

def convert_doc_to_pdf(input_path, output_path):

    """Convert a .docx file to PDF."""

    pdf = PDF()

    pdf.add_page()

    pdf.set_font("Arial", size=12)



    doc = Document(input_path)

    for paragraph in doc.paragraphs:

        pdf.multi_cell(0, 10, sanitize_text(paragraph.text))



    pdf.output(output_path)



def convert_pdf_to_text(input_path):

    """Convert a PDF file to text."""

    return extract_text(input_path)



def convert_pdf_to_docx(input_path, output_path):

    """Convert a PDF file to .docx by extracting text and saving to DOCX."""

    text = convert_pdf_to_text(input_path)

    doc = Document()

    for line in text.splitlines():

        doc.add_paragraph(sanitize_text(line))

    doc.save(output_path)



def convert_heif_to_jpg(input_path, output_path):

    """Convert HEIF file to JPG."""

    img = Image.open(input_path)

    img.save(output_path, "JPEG")



def convert_heif_to_png(input_path, output_path):

    """Convert HEIF file to PNG."""

    img = Image.open(input_path)

    img.save(output_path, "PNG")



def convert_jpg_to_heif(input_path, output_path):

    """Convert JPG file to HEIF."""

    img = Image.open(input_path)

    img.save(output_path, "HEIF")



def convert_png_to_heif(input_path, output_path):

    """Convert PNG file to HEIF."""

    img = Image.open(input_path)

    img.save(output_path, "HEIF")



def convert_jpg_to_png(input_path, output_path):

    """Convert JPG file to PNG."""

    img = Image.open(input_path)

    img.save(output_path, "PNG")



def convert_png_to_jpg(input_path, output_path):

    """Convert PNG file to JPG."""

    img = Image.open(input_path)

    img.save(output_path, "JPEG")



def convert_text_to_pdf(text, output_path):

    """Convert text input to a PDF file."""

    pdf = PDF()

    pdf.add_page()

    pdf.set_font("Arial", size=12)

    pdf.multi_cell(0, 10, sanitize_text(text))

    pdf.output(output_path)



# File Selection Functions

def select_input_file(var):

    """Open a file dialog to select the input file."""

    file_path = filedialog.askopenfilename(filetypes=[("All Files", "*.*")])

    if file_path:

        var.set(file_path)



def select_output_file(var, ext):

    """Open a file dialog to select the output file."""

    file_path = filedialog.asksaveasfilename(defaultextension=ext, filetypes=[(f"{ext.upper()} files", f"*{ext}")])

    if file_path:

        var.set(file_path)



# Dropdown-based Conversion Logic

def on_convert():

    """Handle file conversion based on selected conversion type."""

    input_file = input_file_var.get()

    output_file = output_file_var.get()

    

    if not input_file or not output_file:

        messagebox.showerror("Error", "Please select input and output files.")

        return

    

    conversion_type = conversion_type_var.get()

    

    try:

        # Document Conversions

        if conversion_type == "DOCX to PDF":

            convert_doc_to_pdf(input_file, output_file)

            messagebox.showinfo("Success", "DOCX converted to PDF successfully.")

        

        elif conversion_type == "PDF to DOCX":

            convert_pdf_to_docx(input_file, output_file)

            messagebox.showinfo("Success", "PDF converted to DOCX successfully.")

        

        elif conversion_type == "PDF to Text":

            text = convert_pdf_to_text(input_file)

            with open(output_file.replace(".pdf", ".txt"), "w") as text_file:

                text_file.write(text)

            messagebox.showinfo("Success", "PDF converted to Text successfully.")



        # Image Conversions

        elif conversion_type == "HEIF to JPG":

            convert_heif_to_jpg(input_file, output_file)

            messagebox.showinfo("Success", "HEIF converted to JPG successfully.")



        elif conversion_type == "HEIF to PNG":

            convert_heif_to_png(input_file, output_file)

            messagebox.showinfo("Success", "HEIF converted to PNG successfully.")

        

        elif conversion_type == "JPG to HEIF":

            convert_jpg_to_heif(input_file, output_file)

            messagebox.showinfo("Success", "JPG converted to HEIF successfully.")



        elif conversion_type == "PNG to HEIF":

            convert_png_to_heif(input_file, output_file)

            messagebox.showinfo("Success", "PNG converted to HEIF successfully.")

        

        elif conversion_type == "JPG to PNG":

            convert_jpg_to_png(input_file, output_file)

            messagebox.showinfo("Success", "JPG converted to PNG successfully.")

        

        elif conversion_type == "PNG to JPG":

            convert_png_to_jpg(input_file, output_file)

            messagebox.showinfo("Success", "PNG converted to JPG successfully.")

        

        elif conversion_type == "Text to PDF":

            text = text_input.get("1.0", END)

            convert_text_to_pdf(text, output_file)

            messagebox.showinfo("Success", "Text converted to PDF successfully.")

        

        else:

            messagebox.showerror("Error", "Unsupported conversion type.")

    

    except Exception as e:

        messagebox.showerror("Error", str(e))



def on_enter(event):

    """Change button color on hover."""

    event.widget['background'] = 'lightblue'



def on_leave(event):

    """Reset button color on leave."""

    event.widget['background'] = 'lightgray'



# Initialize the main window

root = Tk()

root.title("Toolz - Document and Image Converter")

root.geometry("600x600")



# Load and set a background image

bg_image = None

try:

    image_path = r"/images/test.png"  # Ensure this path is correct

    if not os.path.exists(image_path):

        raise FileNotFoundError(f"The image file does not exist at {image_path}")

    

    background_image = Image.open(image_path)

    background_image = background_image.resize((600, 600), Image.LANCZOS)

    bg_image = ImageTk.PhotoImage(background_image)



    bg_label = Label(root, image=bg_image)

    bg_label.place(relwidth=1, relheight=1)

except FileNotFoundError as e:

    messagebox.showerror("Error", str(e))

    root.quit()  # Exit if the image cannot be loaded

except OSError as e:

    messagebox.showerror("Error", f"Failed to load image: {e}")

    root.quit()  # Exit if the image cannot be loaded



# Variables to hold file paths and conversion type

input_file_var = StringVar()

output_file_var = StringVar()

conversion_type_var = StringVar(value="Select Conversion Type")



# Conversion Type Dropdown

conversion_types = [

    "DOCX to PDF", "PDF to DOCX", "PDF to Text",

    "HEIF to JPG", "HEIF to PNG", "JPG to HEIF",

    "PNG to HEIF", "JPG to PNG", "PNG to JPG", "Text to PDF"

]



Label(root, text="Select Conversion Type:", bg='white').pack(pady=10)

conversion_type_menu = OptionMenu(root, conversion_type_var, *conversion_types)

conversion_type_menu.pack(pady=10)



# Input and Output File Selection

Label(root, text="Input File: (For Document/Image)").pack(pady=5)

input_entry = Entry(root, textvariable=input_file_var, width=50, bg='lightyellow')

input_entry.pack(pady=5)

Button(root, text="Browse", command=lambda: select_input_file(input_file_var)).pack(pady=5)



Label(root, text="Output File:").pack(pady=5)

output_entry = Entry(root, textvariable=output_file_var, width=50, bg='lightyellow')

output_entry.pack(pady=5)

Button(root, text="Browse", command=lambda: select_output_file(output_file_var, ".pdf")).pack(pady=5)



# Text Input Box for Text to PDF conversion

Label(root, text="Enter Text to Convert to PDF:", bg='white').pack(pady=5)

text_input = Text(root, width=60, height=10, bg='lightyellow')

text_input.pack(pady=5)



# Convert Button

convert_button = Button(root, text="Convert", command=on_convert)

convert_button.pack(pady=20)

convert_button.bind("<Enter>", on_enter)

convert_button.bind("<Leave>", on_leave)



# Run the application

root.mainloop()

